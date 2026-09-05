from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    return h

def add_bold_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

def add_para(text):
    return doc.add_paragraph(text)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    return table

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_mixed_para(parts):
    p = doc.add_paragraph()
    for text, bold in parts:
        run = p.add_run(text)
        run.bold = bold
    return p

# ========== BUILD DOCUMENT ==========

title = doc.add_heading('UPLIFT Mobile — Final Deliverables Report', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

add_para('MIT App Inventor Mobile Application — Backend: UPLIFT AI-Powered PWD Job Matching System')
add_para('Submission Date: June 13, 2026')
doc.add_paragraph()

# ===== PART I =====
add_heading('PART I: TESTING & DEBUGGING REPORT', level=1)

add_heading('1.1 Test Plan', level=2)

add_table(
    ['Test ID', 'Screen', 'Test Case', 'Expected Result', 'Status'],
    [
        ['TC-01', 'Splash', 'Launch with valid token in TinyDB', 'Auto-skip to Dashboard', '✅ Fixed'],
        ['TC-02', 'Splash', 'Launch with expired/invalid token', 'Clear token, show Login', '✅ Fixed'],
        ['TC-03', 'Splash', 'Launch with no token', 'Show Login screen', '✅ Pass'],
        ['TC-04', 'Login', 'Submit with empty email', 'Show email error', '✅ Pass'],
        ['TC-05', 'Login', 'Submit with empty password', 'Show password error', '✅ Pass'],
        ['TC-06', 'Login', 'Submit with invalid email format', 'Show email format error', '✅ Pass'],
        ['TC-07', 'Login', 'Submit with short password (<6 chars)', 'Show password length error', '✅ Pass'],
        ['TC-08', 'Login', 'Submit valid credentials', 'Store token, navigate to Dashboard', '✅ Pass'],
        ['TC-09', 'Login', 'Submit wrong credentials', 'Show 401 error', '✅ Pass'],
        ['TC-10', 'Register', 'Submit with empty name', 'Show error', '✅ Pass'],
        ['TC-11', 'Register', 'Submit with mismatched passwords', 'Show error', '✅ Fixed'],
        ['TC-12', 'Register', 'Submit with existing email', 'Show already registered', '✅ Pass'],
        ['TC-13', 'Register', 'Submit valid registration', 'Success notifier, switch to login', '✅ Pass'],
        ['TC-14', 'Dashboard', 'Load with valid token', 'Display user name and profile', '✅ Pass'],
        ['TC-15', 'Dashboard', 'Token expired / 401 from /me', 'Session expired, redirect to login', '✅ Fixed'],
        ['TC-16', 'Dashboard', 'Navigate to all screens', 'Each button opens correct screen', '✅ Pass'],
        ['TC-17', 'JobSearch', 'Submit empty search query', 'Show "Please enter a search term"', '✅ Pass'],
        ['TC-18', 'JobSearch', 'Valid search with results', 'Display jobs in ListView', '✅ Pass'],
        ['TC-19', 'JobSearch', 'Valid search with no results', 'Show "No results" message', '✅ Pass'],
        ['TC-20', 'JobSearch', 'Server offline during search', 'Show "Match failed" error', '✅ Fixed'],
        ['TC-21', 'JobSearch', 'Adjust sliders', 'Labels update in real-time', '✅ Pass'],
        ['TC-22', 'JobSearch', 'Sliders persist after app restart', 'Restore from TinyDB', '✅ Pass'],
        ['TC-23', 'JobSearch', 'Tap job in ListView', 'Navigate to JobDetails', '✅ Pass'],
        ['TC-24', 'JobSearch', 'Tap "Save Job"', 'Added to TinyDB saved_jobs', '✅ Pass'],
        ['TC-25', 'JobDetails', 'Load job details', 'Display title, scores, AI report', '✅ Pass'],
        ['TC-26', 'JobDetails', 'Apply for job', 'Success notification', '✅ Pass'],
        ['TC-27', 'JobDetails', 'Apply when server unreachable', 'Show error notifier', '✅ Fixed'],
        ['TC-28', 'JobDetails', 'Back to Search', 'Navigate back', '✅ Pass'],
        ['TC-29', 'SavedJobs', 'Load with saved jobs', 'Display saved list', '✅ Pass'],
        ['TC-30', 'SavedJobs', 'Load with no saved jobs', 'Show empty state', '✅ Pass'],
        ['TC-31', 'SavedJobs', 'Remove individual job', 'Job removed from TinyDB', '✅ Pass'],
        ['TC-32', 'SavedJobs', 'Clear all saved jobs', 'All removed, empty state', '✅ Pass'],
        ['TC-33', 'SavedJobs', 'Tap saved job', 'Navigate to JobDetails', '✅ Pass'],
        ['TC-34', 'Applications', 'Load applications list', 'Display from backend', '✅ Pass'],
        ['TC-35', 'Applications', 'No applications yet', 'Show empty state', '✅ Pass'],
        ['TC-36', 'Profile', 'Load with saved data', 'Fields populated from TinyDB', '✅ Pass'],
        ['TC-37', 'Profile', 'Save with empty skills', 'Show "Enter at least one skill"', '✅ Pass'],
        ['TC-38', 'Profile', 'Save valid profile', 'Backend + TinyDB updated', '✅ Pass'],
        ['TC-39', 'Profile', 'Clear local data', 'All profile tags cleared', '✅ Pass'],
        ['TC-40', 'Draft', 'Save with empty cover letter', 'Show "Write something first"', '✅ Pass'],
        ['TC-41', 'Draft', 'Save valid draft', 'Stored in TinyDB with timestamp', '✅ Pass'],
        ['TC-42', 'Draft', 'Load existing draft', 'Text and timestamp restored', '✅ Pass'],
        ['TC-43', 'Draft', 'Load non-existent draft', 'Show "No draft for this job"', '✅ Pass'],
        ['TC-44', 'Draft', 'Delete draft', 'Removed from TinyDB', '✅ Pass'],
        ['TC-45', 'Global', 'Logout', 'Token cleared, return to Login', '✅ Pass'],
        ['TC-46', 'Global', 'App close and reopen', 'Token + saved jobs persist', '✅ Pass'],
        ['TC-47', 'Global', 'Network timeout handling', 'Notifier shows error', '✅ Fixed'],
        ['TC-48', 'Global', 'Rapid button tapping', 'Button disabled during API call', '✅ Fixed'],
    ]
)

doc.add_paragraph()
add_heading('1.2 Bugs Found and Fixes', level=2)

add_table(
    ['Bug ID', 'Screen', 'Issue', 'Root Cause', 'Fix Applied'],
    [
        ['BUG-01', 'Splash', 'Crash when token is expired', 'No 401 handler for /api/auth/me', 'Added GotText handler: clear token, redirect to Login'],
        ['BUG-02', 'Login', 'Register fields not cleared after success', 'No field reset on toggle back', 'Set all register fields to "" on toggle'],
        ['BUG-03', 'Login', 'Button stays disabled after API error', 'Enabled flag never restored', 'Re-enable button in error branch'],
        ['BUG-04', 'Login', 'Double-tap submits multiple API calls', 'No guard against rapid clicking', 'Disable button at start, re-enable on error'],
        ['BUG-05', 'JobSearch', 'ListView shows "null" for missing score', 'JSON path failure', 'Fallback: use "N/A" if score missing'],
        ['BUG-06', 'JobSearch', 'Slider resets to 0.5 despite saved 0.8', 'Empty string used as slider value', 'Use OR 0.5 default if GetValue empty'],
        ['BUG-07', 'JobSearch', 'Save Job duplicates same entry', 'No duplicate ID check', 'Added existence check before adding'],
        ['BUG-08', 'JobDetails', 'Scores blank from SavedJobs nav', 'selected_job_json not set', 'Added StoreValue in SavedJobs AfterPicking'],
        ['BUG-09', 'JobDetails', 'Apply shows success on server error', 'Only 200 checked', 'Added error branch, restore button state'],
        ['BUG-10', 'SavedJobs', 'Remove does nothing if no selection', 'No selection check', 'Added "Please select a job first" guard'],
        ['BUG-11', 'SavedJobs', 'ListView not refreshing after delete', 'Screen not re-initialized', 'Re-run Initialize after delete'],
        ['BUG-12', 'Draft', 'Loads wrong draft after job switch', 'Stale selected_job_id', 'Added confirmation dialog before loading'],
        ['BUG-13', 'Global', 'TinyDB accumulates stale data', 'No cleanup on logout', 'Clear all user tags on logout'],
        ['BUG-14', 'Global', 'GotText updates wrong screen after nav', 'Async response after navigation', 'Check Screen.Visible before UI updates'],
        ['BUG-15', 'Applications', 'Crash if job_title is null', 'Null from backend for deleted jobs', 'Use "Unknown Job" fallback'],
    ]
)

doc.add_paragraph()
add_heading('1.3 Test User Feedback', level=2)

add_bold_para('Tester 1 — Classmate (non-technical):')
add_para('"The login and registration was easy. I liked that it told me exactly what I did wrong instead of just saying error. The job search gave me results that actually matched my profile. One thing — I accidentally tapped the same job twice and it tried to apply twice. Maybe add an Already applied check?"')
add_bold_para('Fix applied: Added duplicate application check and client-side guard against double-submit.')

add_bold_para('Tester 2 — Classmate (technical):')
add_para('"The app worked well overall. I noticed that when I saved a job from search results and then went to Saved Jobs, the list showed the job title but I could not tell which employer it was from. Maybe show both in the saved list."')
add_bold_para('Fix applied: Changed SavedJobs ListView to "Job Title @ Employer Name" format.')

add_bold_para('Tester 3 — Friend with visual impairment (low vision):')
add_para('"The colors were good — I could read the text easily. The buttons were big enough to tap. But the slider controls were hard to use because they are small. Also, I wish the app could read the job details out loud."')
add_bold_para('Fix applied: Increased slider numeric label visibility. Noted TTS as top future priority.')

add_bold_para('Tester 4 — Instructor feedback:')
add_para('"Make sure the app behaves well when the server is down. Also check that the back button on Android does not break navigation."')
add_bold_para('Fix applied: Added network timeout handling and Screen.BackPressed handlers on all screens.')

add_heading('1.4 What Still Needs Improvement', level=2)
improvements = [
    ('Text-to-Speech integration — ', 'For users with visual impairments, automatic audio reading of job details and AI reports.'),
    ('Slider UX — ', 'MIT App Inventor slider is small. Custom numeric input or +/- buttons would be more accessible.'),
    ('Offline job caching — ', 'Saved jobs are cached but full search requires internet. Cache recent searches.'),
    ('Push notifications — ', 'Notify users when application status changes. Requires server-side WebSocket or polling.'),
    ('Multi-language support — ', 'Tagalog option for Filipino users.'),
    ('Biometric login — ', 'Fingerprint or face recognition for faster re-authentication.'),
]
for bold_part, text in improvements:
    add_bullet(text, bold_prefix=bold_part)

add_heading('1.5 Screenshots of Fixed Parts', level=2)
add_para('[Insert screenshots showing: Login error label after fix, JobDetails scores loading from SavedJobs, SavedJobs "Job @ Employer" format, duplicate application guard, "Please select a job first" error]')

doc.add_page_break()

# ===== PART II =====
add_heading('PART II: FINAL UI POLISH & BRANDING', level=1)

add_heading('2.1 What Changed from the Original UI', level=2)

add_table(
    ['Aspect', 'Original Design', 'Final Design', 'Rationale'],
    [
        ['Splash screen', 'Basic text + button', 'Logo, gradient background, subtle animation', 'Professional first impression'],
        ['Button styling', 'Flat rectangles', 'Corner radius 10px, consistent 50px height, disabled mute color', 'Softer modern look, clear disabled state'],
        ['Color palette', 'Ad-hoc hex values', '8-color system: Primary, Success, Warning, Danger, Dark, Light, Gray, Muted', 'Consistency across all screens'],
        ['Typography', 'Mixed font sizes', 'Headers=24px Bold, Sub=18px, Body=16px, Labels=14px', 'Clear readability hierarchy'],
        ['Login screen', 'Single form with toggle', 'Tab-style Login/Register tabs, no layout shift', 'Clearer UX mode indicator'],
        ['Dashboard', 'Simple button list', 'Card sections: Profile, Quick Actions, Account', 'Visual grouping for scannability'],
        ['Job search results', 'Plain text ListView', 'Bold title, gray employer, score badge on right', 'Quick info identification'],
        ['Job details', 'Plain scroll', '2x2 score card grid, bordered report card, red barriers card', 'Scannable AI report layout'],
        ['Saved jobs', 'Minimal', 'Emoji indicators, empty state, confirmation dialogs', 'Polished, prevents data loss'],
        ['Profile screen', 'Basic text areas', 'Section labels, character hints, save feedback', 'Intuitive form layout'],
        ['Loading states', 'None', 'Labels + button text changes during API calls', 'User knows app is working'],
        ['Error displays', 'Inline red label only', 'Inline for forms, Notifier for API, status for loading', 'Right tool per situation'],
        ['App icon', 'Default MIT icon', 'Custom UPLIFT icon: white "U" on dark blue circle', 'Brand identity on home screen'],
        ['Screen transitions', 'Instant', 'Fade animation (300ms) where supported', 'Smoother feel'],
    ]
)

add_heading('2.2 Why We Made Those Changes', level=2)

reasons = [
    'Accessibility-first redesign: The original design was functional but did not fully meet PWD accessibility guidelines. We formalized the color system, increased contrast ratios, standardized font sizes, and added touch target minimums. The tab-style login/register prevents confusion — users with cognitive disabilities may not understand toggle visibility changes.',
    'Professional branding: The UPLIFT app represents an AI-powered system. A polished UI creates trust. The custom app icon, consistent color palette, and card-based layout communicate professionalism. The gradient splash screen and smooth transitions make the app feel native rather than a prototype.',
    'User feedback drove changes: Tester 3 (low vision) confirmed the high-contrast colors worked but struggled with sliders. We kept the slider size limit but added the numeric label right next to it so users can see the exact value. Tester 1 accidental double-tap led to disabled-state buttons during API calls.',
    'Information hierarchy: The card-based layout on Dashboard and JobDetails helps users with cognitive disabilities process information in digestible chunks. Each card has a clear header, so users do not have to read everything to find what they need.'
]
for r in reasons:
    add_bullet(r)

add_heading('2.3 What We Learned About Design/Branding', level=2)

learnings = [
    ('1. MIT App Inventor has real design constraints. ',
     'You cannot customize component borders, shadows, or spacing precisely. We worked within these limits by using VerticalArrangements with BackgroundColor to simulate cards, and horizontal arrangements with percentage widths for layout control. Understanding the tools limits was as important as knowing its features.'),
    ('2. Accessibility and aesthetics are not in conflict. ',
     'High-contrast colors, large text, and big touch targets can look modern and professional. The dark slate + blue + emerald palette is both PWD-friendly and visually appealing. The key was choosing colors that work for both — not treating accessibility as a separate checkbox.'),
    ('3. Consistency is harder than creativity. ',
     'Maintaining the same spacing, font sizes, and color values across 9 screens was difficult. MIT App Inventor property editor requires manual entry for each component. We created a Design System Reference cheat sheet and referred to it constantly. Having all hex codes and font sizes documented saved hours of rework.'),
    ('4. User testing reveals what you cannot see. ',
     'We thought the original toggle login/register was clever. Test users found it confusing. We thought the plain ListView was fine. Testers found it hard to scan. Every change we made was validated by someone actually using the app — not by our assumptions.'),
    ('5. Branding is about trust, not just looks. ',
     'UPLIFT handles sensitive data — disabilities, job applications, personal profiles. The polished UI signals that we take data seriously. The consistent design language says "this app is reliable." For PWD users who have been underserved by technology, that trust is critical.'),
]
for bold_part, text in learnings:
    add_bullet(text, bold_prefix=bold_part)

add_heading('2.4 Screenshots of Final UI', level=2)
add_para('[Insert screenshots: Final splash, Login/Register tabs, Dashboard cards, Job search formatted results, Job details score grid, Saved jobs empty state, Profile form, Draft screen status]')

doc.add_page_break()

# ===== PART III =====
add_heading('PART III: FINAL TESTING & SUBMISSION PREP', level=1)

add_heading('3.1 Final App Overview', level=2)

add_mixed_para([('App Name: ', True), ('UPLIFT Mobile', False)])
add_mixed_para([('Purpose: ', True), ('AI-powered job matching platform for Persons with Disabilities (PWDs)', False)])
add_mixed_para([('Platform: ', True), ('MIT App Inventor (Android)', False)])
add_mixed_para([('Backend: ', True), ('UPLIFT — Python FastAPI server with PostgreSQL, pgvector, SentenceTransformer, CrossEncoder, Flan-T5', False)])
add_mixed_para([('Total Screens: ', True), ('9', False)])

add_bold_para('User Flow:')
add_para('Splash -> Login/Register -> Dashboard -> Job Search -> Job Details (Apply) -> Saved Jobs <-> Draft App -> Applications <-> Profile')

add_heading('3.2 Features Completed', level=2)

add_table(
    ['Feature', 'Status', 'Details'],
    [
        ['User Authentication', 'Complete', 'Login with JWT, registration, session persistence via TinyDB, auto-login'],
        ['Profile Management', 'Complete', 'View/edit summary, skills, disabilities; save to backend + cache locally'],
        ['AI Job Matching', 'Complete', 'Semantic search, cross-encoder re-ranking, weighted scoring'],
        ['Job Details & AI Report', 'Complete', 'Safety/Skill/Stamina scores, Flan-T5 analysis, barrier warnings'],
        ['Job Application', 'Complete', 'One-tap apply with auto-generated resume; duplicate detection'],
        ['Saved Jobs (Favorites)', 'Complete', 'Offline-capable; save, view, remove, clear all'],
        ['Application Drafts', 'Complete', 'Per-job notes; save, load, delete; persists in TinyDB'],
        ['Application History', 'Complete', 'View all submitted applications with status'],
        ['Match Weight Customization', 'Complete', 'Safety/Skill/Stamina sliders persisted in TinyDB'],
        ['Error Checking', 'Complete', 'Input validation, API error handling, network timeout'],
        ['Accessibility Design', 'Complete', 'High-contrast, large fonts (16-24px), 50px+ touch targets'],
        ['Logout', 'Complete', 'Clears TinyDB, returns to login'],
    ]
)

add_heading('3.3 Main Challenges and What We Learned', level=2)

add_bold_para('Challenge 1: MIT App Inventor JSON handling is extremely limited.')
add_para('App Inventor has no support for nested JSON access (e.g., obj.parent.child). The UPLIFT backend returns deeply nested match data.')
add_mixed_para([('Solution: ', True), ('Stored full job objects as JSON strings in TinyDB and used chained json.GetValue / json.GetValueFrom calls. Maintained parallel global variable lists alongside TinyDB data.', False)])
add_mixed_para([('Lesson: ', True), ('Flatten your JSON on the backend if possible, or accept creative workarounds for nested data.', False)])

add_bold_para('Challenge 2: Cross-screen state management.')
add_para('MIT App Inventor open another screen only passes a single text start value — not complex objects.')
add_mixed_para([('Solution: ', True), ('Used TinyDB as a shared global state store. Screen A stores with StoreValue, Screen B reads in Initialize. More robust than start values.', False)])
add_mixed_para([('Lesson: ', True), ('TinyDB is the best cross-screen communication mechanism in App Inventor. Design data flow around it from the start.', False)])

add_bold_para('Challenge 3: Asynchronous API responses fighting with navigation.')
add_para('Web.GotText fires asynchronously — the user may navigate away before the response arrives.')
add_mixed_para([('Solution: ', True), ('Added a screen_active flag stored in TinyDB. In GotText, check this flag before any UI update. Non-critical errors use Notifier.ShowAlert.', False)])
add_mixed_para([('Lesson: ', True), ('Always assume the user will navigate faster than your API responds. Guard every asynchronous callback.', False)])

add_bold_para('Challenge 4: Working within MIT App Inventor component limitations.')
add_para('No borders, shadows, or precise spacing control. Slider is small. ListView has minimal formatting.')
add_mixed_para([('Solution: ', True), ('Used arrangements with BackgroundColor to simulate cards. Invisible labels for spacing. FontBold/FontSize for visual hierarchy. Real-time numeric label next to sliders.', False)])
add_mixed_para([('Lesson: ', True), ('Great App Inventor apps work with the platforms strengths and use creative tricks for limitations.', False)])

add_bold_para('Challenge 5: Making PWD accessibility work within App Inventor.')
add_para('Limited built-in accessibility support — no screen reader optimization, no focus management.')
add_mixed_para([('Solution: ', True), ('Focused on what we could control: high contrast ratios (verified with checker), large touch targets, clear error messages, consistent navigation, readable fonts. Documented gaps as future work.', False)])
add_mixed_para([('Lesson: ', True), ('Accessibility is a spectrum. Making the app better for PWDs — even without perfect tooling — is still valuable. Design system changes alone significantly improved usability.', False)])

add_heading('3.4 Blocks Organization', level=2)
blocks_rules = [
    ('Color coding: ', 'Purple = Screen init, Blue = Button click, Green = Web response, Red = Error check, Orange = TinyDB ops'),
    ('Comments: ', 'Each major section has a label: --- ERROR CHECKING ---, --- API CALL ---, --- STORE IN TINYDB ---'),
    ('Grouping: ', 'Related blocks adjacent on workspace. Each screen in separate .bky export.'),
    ('Unused blocks: ', 'All test blocks removed. Final .aia has zero unused components.'),
    ('Naming: ', 'Consistent prefix: Button_ActionName, TextBox_FieldName, Label_Description, Web_EndpointName'),
]
for bold_part, text in blocks_rules:
    add_bullet(text, bold_prefix=bold_part)

add_heading('3.5 Final File Checklist', level=2)
checklist = [
    'UPLIFT_Mobile_Final.aia — MIT App Inventor project file',
    'All 9 screens created and labeled',
    'All block comments added',
    'All unused components removed',
    'No test/debug blocks remaining',
    'TinyDB tag naming consistent across all screens',
    'api_base_url configurable (not hardcoded)',
    'Android back button behavior defined per screen',
    'Screenshots captured (see checklist below)',
    'Final summary document prepared',
    'Presentation slides created (see presentation_template_30min.md)',
]
for item in checklist:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run('[ ] ')
    run.bold = False
    p.add_run(item)

add_heading('3.6 Final Screenshot Checklist', level=2)
add_table(
    ['#', 'Screen', 'Content'],
    [
        ['1', 'Splash', 'Branded splash with UPLIFT logo + "Get Started"'],
        ['2', 'Login', 'Tab-style with email filled, Login button'],
        ['3', 'Register', 'Tab toggled to Register, all fields filled'],
        ['4', 'Dashboard', 'Welcome message, profile card, nav buttons'],
        ['5', 'JobSearch', 'Sliders adjusted, search results visible'],
        ['6', 'JobDetails', 'Score cards (2x2 grid), AI report, Apply button'],
        ['7', 'SavedJobs', 'List of saved jobs with Remove + Clear All'],
        ['8', 'Draft', 'Cover letter textbox with Save/Load/Delete buttons'],
        ['9', 'Applications', 'List of submitted applications with statuses'],
        ['10', 'Profile', 'All fields populated, Save button visible'],
        ['11', 'Error state', 'Login with empty email showing red error label'],
        ['12', 'Error state', 'Search with empty query showing Notifier dialog'],
        ['13', 'Persistence', 'Dashboard after app restart — same user data'],
    ]
)

add_heading('3.7 Group Presentation Preparation', level=2)
add_para('Use the companion file presentation_template_30min.md for the final group presentation. It contains 15 ready-to-use slides with speaker notes, timing breakdown (30-35 minutes), image/screenshot placement instructions, and Q&A section.')

add_bold_para('Key talking points for the presentation:')
presentation_points = [
    ('Relevance: ', '28% PWD employment rate — UPLIFT exists to bridge this gap with AI'),
    ('System integration: ', 'MIT App Inventor frontend <-> FastAPI backend with AI pipeline'),
    ('Accessibility first: ', 'Design decisions rooted in PWD needs, not afterthought'),
    ('Uniqueness: ', 'Multi-stage AI (not just keywords) + fairness auditing + PWD-first design'),
    ('Enhancement: ', 'Offline favorites and drafts via TinyDB'),
    ('Challenges overcome: ', 'JSON parsing, cross-screen state, async navigation, MIT constraints'),
]
for bold_part, text in presentation_points:
    add_bullet(text, bold_prefix=bold_part)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('--- End of Final Deliverables Report ---')
run.bold = True
run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

doc.save(r'C:\Users\reyma\Desktop\UPLIFT-Thesis\UPLIFT_Mobile_Final_Deliverables.docx')
print("DOCX created successfully!")
